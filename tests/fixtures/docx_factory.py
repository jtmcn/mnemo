"""Factory functions for creating test DOCX files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def create_test_docx(
    output_path: Path,
    title: str = "Test DOCX Book",
    author: str = "Test Author",
) -> Path:
    """Create a minimal DOCX with headings and paragraphs."""
    doc = Document()
    doc.core_properties.title = title
    doc.core_properties.author = author

    doc.add_heading("Chapter 1: Basics", level=1)
    doc.add_paragraph("This chapter covers the basic concepts.")
    doc.add_heading("Section 1.1: Setup", level=2)
    doc.add_paragraph("Setup instructions go here.")
    doc.add_heading("Chapter 2: Advanced", level=1)
    doc.add_paragraph("Advanced topics are covered in this chapter.")

    doc.save(output_path)
    return output_path


def create_docx_with_code(output_path: Path) -> Path:
    """Create a DOCX with code blocks, tables, and mixed content."""
    doc = Document()
    doc.core_properties.title = "Code Examples"
    doc.core_properties.author = "Code Author"

    doc.add_heading("Chapter 1: Code", level=1)
    doc.add_paragraph("Here is a code example:")

    # Add a code-style paragraph
    code_para = doc.add_paragraph("def hello():\n    print('hello')")
    code_para.style = doc.styles["Normal"]
    # Set font to monospace to trigger code detection
    for run in code_para.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    doc.add_paragraph("And some more text after the code.")

    # Add a table
    doc.add_heading("Section 1.1: Data", level=2)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "alpha"
    table.cell(1, 1).text = "1"
    table.cell(2, 0).text = "beta"
    table.cell(2, 1).text = "2"

    doc.save(output_path)
    return output_path


def create_docx_no_headings(output_path: Path) -> Path:
    """Create a DOCX with no headings (structure_source should be 'inferred')."""
    doc = Document()
    doc.core_properties.title = "Flat Document"
    doc.core_properties.author = "Flat Author"

    doc.add_paragraph("First paragraph of text.")
    doc.add_paragraph("Second paragraph of text.")
    doc.add_paragraph("Third paragraph of text.")

    doc.save(output_path)
    return output_path


def create_docx_no_metadata(output_path: Path) -> Path:
    """Create a DOCX with no title or author metadata."""
    doc = Document()
    # Don't set title or author
    doc.add_heading("Some Content", level=1)
    doc.add_paragraph("Text here.")
    doc.save(output_path)
    return output_path


def create_docx_multi_author(output_path: Path) -> Path:
    """Create a DOCX with multiple authors (semicolon-separated)."""
    doc = Document()
    doc.core_properties.title = "Multi Author Book"
    doc.core_properties.author = "Alice Smith; Bob Jones; Carol White"
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("Content.")
    doc.save(output_path)
    return output_path
