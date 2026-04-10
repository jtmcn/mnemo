"""Tests for the DOCX parser."""

from __future__ import annotations

from pathlib import Path

import pytest

from mnemo.docx import DocxParser
from mnemo.models import ContentType
from tests.fixtures.docx_factory import (
    create_docx_multi_author,
    create_docx_no_headings,
    create_docx_no_metadata,
    create_docx_with_code,
    create_test_docx,
)


@pytest.fixture
def parser() -> DocxParser:
    return DocxParser()


class TestDocxMetadata:
    """Tests for metadata extraction from DOCX."""

    def test_extracts_title_and_author(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_test_docx(tmp_path / "test.docx", title="My Book", author="Author One")
        book, _ = parser.parse(docx_path)
        assert book.title == "My Book"
        assert book.authors == ["Author One"]

    def test_multi_author_semicolons(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_docx_multi_author(tmp_path / "multi.docx")
        book, _ = parser.parse(docx_path)
        assert book.title == "Multi Author Book"
        assert book.authors == ["Alice Smith", "Bob Jones", "Carol White"]

    def test_no_metadata_uses_filename(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_docx_no_metadata(tmp_path / "untitled.docx")
        book, _ = parser.parse(docx_path)
        assert book.title == "untitled"

    def test_generates_book_id(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_test_docx(tmp_path / "test.docx")
        book, _ = parser.parse(docx_path)
        assert len(book.id) == 6
        assert all(c in "0123456789abcdef" for c in book.id)

    def test_generates_file_hash(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_test_docx(tmp_path / "test.docx")
        book, _ = parser.parse(docx_path)
        assert len(book.file_hash) == 64

    def test_file_not_found(self, parser: DocxParser) -> None:
        with pytest.raises(FileNotFoundError):
            parser.parse(Path("/nonexistent/file.docx"))


class TestDocxContent:
    """Tests for content extraction from DOCX."""

    def test_heading_creates_section_path(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_test_docx(tmp_path / "test.docx")
        _, blocks = parser.parse(docx_path)

        # Find blocks under Chapter 1
        ch1_blocks = [b for b in blocks if "Chapter 1: Basics" in b.section_path]
        assert len(ch1_blocks) > 0

        # Find blocks under Section 1.1 (nested)
        sec_blocks = [b for b in blocks if "Section 1.1: Setup" in b.section_path]
        assert len(sec_blocks) > 0
        # Section 1.1 should be nested under Chapter 1
        for b in sec_blocks:
            assert "Chapter 1: Basics" in b.section_path

    def test_heading_level_resets_stack(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_test_docx(tmp_path / "test.docx")
        _, blocks = parser.parse(docx_path)

        # Chapter 2 blocks should NOT have Chapter 1 in path
        ch2_blocks = [b for b in blocks if "Chapter 2: Advanced" in b.section_path]
        assert len(ch2_blocks) > 0
        for b in ch2_blocks:
            assert "Chapter 1: Basics" not in b.section_path

    def test_code_detection_by_font(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_docx_with_code(tmp_path / "code.docx")
        _, blocks = parser.parse(docx_path)

        code_blocks = [b for b in blocks if b.content_type == ContentType.CODE]
        assert len(code_blocks) >= 1
        assert "def hello" in code_blocks[0].content

    def test_table_extraction(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_docx_with_code(tmp_path / "code.docx")
        _, blocks = parser.parse(docx_path)

        table_blocks = [b for b in blocks if b.content_type == ContentType.TABLE]
        assert len(table_blocks) == 1
        assert "Name" in table_blocks[0].content
        assert "alpha" in table_blocks[0].content
        assert "|" in table_blocks[0].content

    def test_no_headings_structure_inferred(self, parser: DocxParser, tmp_path: Path) -> None:
        docx_path = create_docx_no_headings(tmp_path / "flat.docx")
        book, blocks = parser.parse(docx_path)
        assert book.structure_source == "inferred"
        assert len(blocks) > 0

    def test_empty_paragraphs_skipped(self, parser: DocxParser, tmp_path: Path) -> None:
        """Empty paragraphs should not create blocks."""
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Sparse"
        doc.add_paragraph("")
        doc.add_paragraph("   ")
        doc.add_paragraph("Real content")
        doc.add_paragraph("")
        path = tmp_path / "sparse.docx"
        doc.save(path)

        _, blocks = parser.parse(path)
        assert len(blocks) == 1
        assert blocks[0].content == "Real content"


class TestDocxAmendmentBoxes:
    """Tests for amendment box detection in single-cell tables."""

    def test_amendment_box_parsed_as_text(self, parser: DocxParser, tmp_path: Path) -> None:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Amendment Doc"
        doc.add_heading("Section 1", level=1)
        doc.add_paragraph("Original paragraph text.")

        # Add a single-cell table mimicking an ERCOT amendment box
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = (
            "[OBDRR046:  Replace the paragraph above with the following "
            "upon system implementation of NPRR1188:]\n"
            "Amended paragraph text with CLR additions."
        )

        path = tmp_path / "amendment.docx"
        doc.save(path)

        _, blocks = parser.parse(path)
        table_blocks = [b for b in blocks if b.content_type == ContentType.TABLE]
        assert len(table_blocks) == 0, "Amendment box should not be a table"

        amendment_blocks = [b for b in blocks if "PENDING AMENDMENT" in b.content]
        assert len(amendment_blocks) == 1
        assert "OBDRR046" in amendment_blocks[0].content
        assert "NPRR1188" in amendment_blocks[0].content
        assert "Amended paragraph text" in amendment_blocks[0].content
        assert amendment_blocks[0].content_type == ContentType.TEXT

    def test_regular_table_not_affected(self, parser: DocxParser, tmp_path: Path) -> None:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Table Doc"
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"

        path = tmp_path / "regular_table.docx"
        doc.save(path)

        _, blocks = parser.parse(path)
        table_blocks = [b for b in blocks if b.content_type == ContentType.TABLE]
        assert len(table_blocks) == 1

    def test_single_cell_non_amendment_stays_table(
        self, parser: DocxParser, tmp_path: Path
    ) -> None:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Note Doc"
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "This is just a note in a box."

        path = tmp_path / "note.docx"
        doc.save(path)

        _, blocks = parser.parse(path)
        table_blocks = [b for b in blocks if b.content_type == ContentType.TABLE]
        assert len(table_blocks) == 1
        assert "PENDING AMENDMENT" not in table_blocks[0].content


class TestDocxLanguageDetection:
    """Tests for code language detection."""

    def test_python_heuristic(self, parser: DocxParser, tmp_path: Path) -> None:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Python Doc"
        doc.add_heading("Code", level=1)
        p = doc.add_paragraph("def main():\n    pass")
        for run in p.runs:
            run.font.name = "Courier New"

        path = tmp_path / "python.docx"
        doc.save(path)

        _, blocks = parser.parse(path)
        code_blocks = [b for b in blocks if b.content_type == ContentType.CODE]
        assert len(code_blocks) == 1
        assert code_blocks[0].language == "python"

    def test_shell_shebang(self, parser: DocxParser, tmp_path: Path) -> None:
        from docx import Document

        doc = Document()
        doc.core_properties.title = "Shell Doc"
        p = doc.add_paragraph("#!/bin/bash\necho hello")
        for run in p.runs:
            run.font.name = "Courier New"

        path = tmp_path / "shell.docx"
        doc.save(path)

        _, blocks = parser.parse(path)
        code_blocks = [b for b in blocks if b.content_type == ContentType.CODE]
        assert len(code_blocks) == 1
        assert code_blocks[0].language == "shell"
